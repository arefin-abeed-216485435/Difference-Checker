import docx
from docx.enum.text import WD_COLOR_INDEX

doc = docx.Document()
doc1= docx.Document()

for line in open('oldConf_path.txt').readlines():
    old = line

for line in open('newConf_path.txt').readlines():
    new = line    

with open(new, 'r') as file1:
    with open(old, 'r') as file2:
        diff = set(file1).difference(file2)
        
with open(new, 'r') as file1:
    with open(old, 'r') as file2:
        diff1 = set(file2).difference(file1)        
         
diff.discard('\n')
diff1.discard('\n')

doc.add_paragraph('The difference(s) from the updated version are: \n')
doc1.add_paragraph('The difference(s) from the old version are: \n')
para = doc.add_paragraph('')
para1 = doc1.add_paragraph('')
with open(new, 'r') as file1:
        for line in file1:
                if line in diff:
                    para.add_run(line).font.highlight_color = WD_COLOR_INDEX.YELLOW       
                else:
                    para.add_run(line)

with open(old, 'r') as file2:
        for line in file2:
                if line in diff1:
                    para1.add_run(line).font.highlight_color = WD_COLOR_INDEX.YELLOW       
                else:
                    para1.add_run(line)  


doc.save("differences_from_updated.docx")
doc1.save("differences_from_old.docx")